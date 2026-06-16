"""Offline tests for Drive extraction + master/voice profile synthesis (no network/LLM)."""
import io
import unittest
from unittest import mock

from job_agent import db, drive
from job_agent.reasoning import master_profile as mp


class FakeFilesApi:
    """Minimal stand-in for svc.files() supporting .list().execute() pagination."""
    def __init__(self, pages):
        self._pages, self._i = pages, 0

    def list(self, **kw):
        return self

    def execute(self):
        page = self._pages[self._i]
        self._i += 1
        return page


class FakeSvc:
    def __init__(self, pages):
        self._api = FakeFilesApi(pages)

    def files(self):
        return self._api


class TestDriveExtraction(unittest.TestCase):
    def test_docx_text_roundtrip(self):
        import docx
        d = docx.Document()
        d.add_paragraph("Muhammad Eltahir")
        d.add_paragraph("Chief of Staff — led strategy and operations.")
        buf = io.BytesIO()
        d.save(buf)
        text = drive._docx_text(buf.getvalue())
        self.assertIn("Muhammad Eltahir", text)
        self.assertIn("Chief of Staff", text)

    def test_find_documents_filters_and_dedupes(self):
        page = {"files": [
            {"id": "1", "name": "Eltahir Resume 2024.pdf", "mimeType": drive.PDF, "modifiedTime": "2024-01-01T00:00:00Z"},
            {"id": "1", "name": "Eltahir Resume 2024.pdf", "mimeType": drive.PDF, "modifiedTime": "2024-01-01T00:00:00Z"},  # dup id
            {"id": "2", "name": "Cover Letter - McKinsey.docx", "mimeType": drive.DOCX, "modifiedTime": "2025-02-01T00:00:00Z"},
            {"id": "3", "name": "Some Folder", "mimeType": "application/vnd.google-apps.folder", "modifiedTime": "2023-01-01T00:00:00Z"},  # excluded
            {"id": "4", "name": "CV old", "mimeType": drive.GOOGLE_DOC, "modifiedTime": "2026-03-01T00:00:00Z"},
        ]}
        files = drive.find_documents(FakeSvc([page]))
        ids = [f["id"] for f in files]
        self.assertEqual(set(ids), {"1", "2", "4"})       # folder excluded, dup collapsed
        self.assertEqual(ids[0], "4")                     # newest first (2026 > 2025 > 2024)

    def test_is_cover_letter(self):
        self.assertTrue(drive.is_cover_letter({"name": "Cover Letter - Stripe.docx"}))
        self.assertFalse(drive.is_cover_letter({"name": "Eltahir Resume.pdf"}))


class TestMasterProfile(unittest.TestCase):
    DOCS = [
        ({"id": "1", "name": "Resume 2024.pdf", "mimeType": drive.PDF, "modifiedTime": "2024-01-01T00:00:00Z"},
         "Corp Dev Manager at Condé Nast. Led M&A."),
        ({"id": "2", "name": "Cover Letter.docx", "mimeType": drive.DOCX, "modifiedTime": "2025-02-01T00:00:00Z"},
         "I am drawn to mission-driven teams..."),
    ]

    def test_docset_hash_changes_with_modified(self):
        a = mp.docset_hash([f for f, _ in self.DOCS])
        b = mp.docset_hash([{"id": "1", "modifiedTime": "2099-01-01T00:00:00Z"},
                            {"id": "2", "modifiedTime": "2025-02-01T00:00:00Z"}])
        self.assertNotEqual(a, b)

    def test_build_master_profile_passes_through_llm(self):
        fake = {"name": "Muhammad Eltahir", "experience_threads": ["Strategy", "Operations"], "employers": []}
        with mock.patch.object(mp.llm, "complete_json", return_value=fake) as m:
            out = mp.build_master_profile(self.DOCS, model="x")
        self.assertEqual(out["name"], "Muhammad Eltahir")
        # both documents were included in the prompt (union, not latest-only)
        prompt = m.call_args.args[0]
        self.assertIn("Resume 2024.pdf", prompt)
        self.assertIn("Cover Letter.docx", prompt)

    def test_voice_profile_no_cover_letters_is_approximate_without_llm(self):
        with mock.patch.object(mp.llm, "complete_json", side_effect=AssertionError("should not call")):
            voice = mp.build_voice_profile([])
        self.assertTrue(voice["approximate"])

    def test_load_or_build_stops_when_nothing_shared(self):
        conn = db.connect(":memory:")
        db.init_db(conn)
        with mock.patch.object(drive, "collect", return_value=([], [])), \
             mock.patch.object(drive, "service_account_email", return_value="sa@x.iam"), \
             mock.patch.object(drive, "build_service", return_value=FakeSvc([{"files": []}])):
            with self.assertRaises(drive.DriveError) as ctx:
                mp.load_or_build(conn)
        self.assertIn("share", str(ctx.exception).lower())
        conn.close()


if __name__ == "__main__":
    unittest.main(verbosity=2)
