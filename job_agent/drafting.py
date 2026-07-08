"""Generate tailored resume + cover-letter drafts for surfaced roles.

Sole source of facts = the cached MASTER PROFILE; tone = the VOICE PROFILE. Output
is editable .md AND .docx under ./applications/<company>-<role>/. Drafts are LOCAL
only (never published) and idempotent (skip already-drafted unless regenerate).

ABSOLUTE GROUNDING: tailoring is selection / emphasis / ordering / rewording of TRUE
content only. The model is instructed never to invent or embellish an employer,
title, date, degree, metric, or skill. If the JD wants something the candidate lacks,
the draft must not claim it (the cover letter may honestly frame transferable work).
"""
from __future__ import annotations

import json
import re
import sqlite3
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from . import config, oauth, store
from .reasoning import llm

def load_profiles() -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """Load the cached master + voice profiles for drafting. Master is required."""
    if not config.MASTER_PROFILE_PATH.exists():
        raise FileNotFoundError("Master profile not built yet — run `job-agent master-profile` first.")
    master = json.loads(config.MASTER_PROFILE_PATH.read_text(encoding="utf-8"))
    voice = (json.loads(config.VOICE_PROFILE_PATH.read_text(encoding="utf-8"))
             if config.VOICE_PROFILE_PATH.exists() else {"approximate": True})
    return master, voice


DRAFT_SYSTEM = (
    "You tailor a candidate's REAL experience to a specific job posting. ABSOLUTE GROUNDING: "
    "you may use ONLY facts that appear in the MASTER PROFILE provided — never invent or "
    "embellish an employer, title, date, degree, metric, certification, or skill. Tailoring "
    "means selecting, emphasizing, ordering, and rewording TRUE content to fit the job; it "
    "NEVER means adding new claims. If the job wants something the candidate does not have, do "
    "NOT claim it — instead the cover letter may honestly frame transferable experience. Write "
    "the cover letter in the candidate's own VOICE (provided). A fabricated credential is a "
    "critical failure. NEVER mention business school, an MBA, or any graduate-school "
    "application, admission, or plan to attend — the candidate has no graduate degree and "
    "does not want this referenced. "
    "STYLE: the cover letter must read like the candidate wrote it himself — plain, confident, "
    "specific prose with varied sentence lengths; contractions are fine. It must NOT read like "
    "AI output: never use em dashes or en dashes anywhere in the cover letter (use commas, "
    "periods, or parentheses), never use stock applicant phrases, and never write three "
    "parallel clauses in a row. Return plain text using the exact ===SECTION=== markers "
    "requested — NOT JSON, no code fences."
)


def _contact_line(master: Dict[str, Any]) -> str:
    c = master.get("contact") or {}
    parts = [c.get("email") or config.NOTIFY_EMAIL, c.get("phone"), c.get("location"), c.get("linkedin")]
    return "  ·  ".join(p for p in parts if p)


def _draft_prompt(master: Dict[str, Any], voice: Dict[str, Any], job: sqlite3.Row) -> str:
    jd = (job["description"] or "")[:6000]
    facts = {k: v for k, v in master.items() if k != "_meta"}  # drop provenance (e.g. filenames)
    return f"""Produce a tailored, ONE-PAGE resume and a cover letter for this candidate and job.

MASTER PROFILE (the ONLY source of facts — do not go beyond it):
{json.dumps(facts, ensure_ascii=False, indent=2)[:16000]}

VOICE PROFILE (imitate this style in the cover letter):
{json.dumps(voice, ensure_ascii=False, indent=2)[:4000]}

JOB:
  Title: {job['title']}
  Company: {job['company']}
  Location: {job['location'] or 'n/a'}
  Description:
  \"\"\"{jd}\"\"\"

Return plain text with these EXACT marker lines (no JSON, no Markdown #/**, no code fences):

===RESUME===
NAME: <full name>
CONTACT: {_contact_line(master)}
SUMMARY: <ONE line, <=24 words, positioning the candidate FOR THIS ROLE — true facts only>
## EXPERIENCE
@ <Company> :: <start – end dates>
> <Title> :: <optional title-specific dates>
- <achievement bullet with the real metric>
- <bullet>
@ <Company> :: <dates>
> <Title>
- <bullet>
## EDUCATION
@ <School> :: <year>
> <Degree>
## SKILLS
- <8–12 skills separated by '  ·  ' on ONE line — the ones THIS job cares about most>

===COVER_LETTER===
<cover letter, 220–300 words, 3–4 paragraphs, addressed to {job['company']}; plain paragraphs separated by blank lines; sign off "Sincerely," then the candidate's name>

===OMITTED===
- <a JD requirement the candidate does NOT clearly meet> (or the single word: none)

CRITICAL rules:
- TAILOR, don't recycle. Study the JD first, then for each recent role pick the 3–4 TRUE
  achievements that best match what THIS job optimizes for, and REWRITE each bullet to lead
  with the outcome this JD cares about — never a near-copy of the profile's wording. Echo the
  JD's own vocabulary where it is truthful. Two different jobs must produce visibly different
  resumes: different bullet selection, different emphasis, different summary.
- ONE PAGE. Keep 10–14 bullets TOTAL: 3–4 for the most relevant/recent roles, 1–2 for older
  ones, and OMIT minor/old roles (brief internships, volunteer/student roles) when needed to
  fit. Skills on ONE line.
- Use the marker structure EXACTLY: NAME:, CONTACT:, SUMMARY:, '## SECTION', '@ Org :: Dates',
  '> Title :: Dates', '- bullet'. One item per line. Do NOT use '#' headings or '**' bold.
- Keep the CONTACT line exactly as given.
- COVER LETTER voice & style (all mandatory):
  * Write in the candidate's own voice per the VOICE PROFILE; vary sentence lengths; be concrete.
  * Name at least one SPECIFIC thing about {job['company']} or this role, drawn from the JD.
  * NO em dashes (—) or en dashes (–) anywhere. Use commas, periods, or parentheses instead.
  * NEVER use these stock phrases (or close variants): "I am excited", "I am writing to
    express", "aligns perfectly", "leverage my", "passionate about", "proven track record",
    "fast-paced environment", "hit the ground running", "uniquely positioned", "resonates",
    "I believe my skills".
  * No more than ONE rhetorical flourish; no lists of three parallel clauses.
- Grounding: every employer, title, date, degree, metric, and skill MUST appear in the master
  profile, with real numbers exactly as written. Never invent. No business-school / MBA mention."""


_SECTION_RE = re.compile(r"^===\s*(RESUME|COVER[_ ]?LETTER|OMITTED)\s*===\s*$", re.I | re.M)
_FENCE_RE = re.compile(r"^```[a-zA-Z0-9]*\n|\n```\s*$")

# --- cover-letter style gate (AI tells the user explicitly flagged) ----------

_BANNED_PHRASES = (
    "i am excited", "i'm excited", "i am writing to express", "aligns perfectly",
    "leverage my", "passionate about", "proven track record", "fast-paced environment",
    "hit the ground running", "uniquely positioned", "resonates", "i believe my skills",
)


def _style_violations(cover: str) -> List[str]:
    """Objections a human reviewer would raise; non-empty triggers one rewrite pass."""
    v = []
    if "—" in cover or "–" in cover:
        v.append("contains em/en dashes (use commas, periods, or parentheses)")
    low = cover.lower()
    v += [f'stock phrase "{p}"' for p in _BANNED_PHRASES if p in low]
    return v


def _scrub_dashes(cover: str) -> str:
    """Last-resort mechanical fix so no dash ever ships even if the retry missed one."""
    cover = re.sub(r"(?<=\d)\s*[–—]\s*(?=\d)", "-", cover)  # numeric ranges -> hyphen
    return re.sub(r"\s*[—–]\s*", ", ", cover)


def _strip_md_bold(s: str) -> str:
    """Drop stray **markers** the model sometimes emits despite instructions."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def _split_sections(text: str):
    """Parse the marker-delimited draft output into (resume_md, cover_md, omitted[])."""
    parts, matches = {}, list(_SECTION_RE.finditer(text))
    for i, m in enumerate(matches):
        key = m.group(1).upper().replace(" ", "_")
        key = "COVER" if key.startswith("COVER") else key
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[m.end():end].strip()
        body = _FENCE_RE.sub("", body).strip()  # tolerate stray code fences
        parts[key] = body
    omitted = [ln.strip(" -*\t") for ln in parts.get("OMITTED", "").splitlines()
               if ln.strip() and ln.strip().lower() not in ("none", "n/a", "- none")]
    return parts.get("RESUME", ""), parts.get("COVER", ""), omitted


# --- markdown -> docx (minimal, dependency-light) ---------------------------

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str) -> None:
    """Render inline **bold** within a docx paragraph; everything else plain."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _build_docx(md: str):
    """Render a small Markdown subset (#/##/### headings, '-'/'*' bullets, **bold**,
    '---' rules, blank-line paragraphs) into a docx.Document."""
    import docx

    doc = docx.Document()
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph().add_run("—" * 20)
        elif line.lstrip().startswith(("- ", "* ")):
            _add_runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:])
        else:
            _add_runs(doc.add_paragraph(), line)
    return doc


def md_to_docx(md: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    _build_docx(md).save(str(path))


def md_to_docx_bytes(md: str) -> bytes:
    import io
    buf = io.BytesIO()
    _build_docx(md).save(buf)
    return buf.getvalue()


# --- résumé / cover-letter renderers (match the candidate's Garamond one-pager) ----

GARAMOND = "Garamond"
_PAGE_W_IN = 8.5


def _tight(p, before=0.0, after=0.0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = _Pt(before)
    pf.space_after = _Pt(after)
    pf.line_spacing = line
    return p


def _run(p, text, size=11.0, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = GARAMOND
    r.font.size = _Pt(size)
    r.bold, r.italic = bold, italic
    return r


def _Pt(v):
    from docx.shared import Pt
    return Pt(v)


def _bottom_border(p):
    """Thin rule under a section header (like the reference resume)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "808080")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_resume_docx(struct: str):
    """Render the structured resume block into a Garamond one-pager matching the
    candidate's format: 22pt centered name, centered contact, 13pt underlined uppercase
    section headers, bold company + right-aligned dates, italic titles, tight bullets."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches

    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = GARAMOND, _Pt(11)
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Inches(0.35), Inches(0.35)
    sec.left_margin, sec.right_margin = Inches(0.7), Inches(0.7)
    content_w = Inches(_PAGE_W_IN - 0.7 - 0.7)

    def right_tab(p):
        p.paragraph_format.tab_stops.add_tab_stop(content_w, WD_TAB_ALIGNMENT.RIGHT)

    for raw in struct.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("NAME:"):
            p = _tight(doc.add_paragraph(), after=2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, line[5:].strip(), size=22)
        elif line.startswith("CONTACT:"):
            p = _tight(doc.add_paragraph(), after=8); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, line[8:].strip(), size=10.5)
        elif line.startswith("SUMMARY:"):
            _run(_tight(doc.add_paragraph(), after=3, line=1.05), line[8:].strip(), size=11, italic=True)
        elif line.startswith("## "):
            p = _tight(doc.add_paragraph(), before=10, after=3)
            _run(p, line[3:].strip().upper(), size=13, bold=True)
            _bottom_border(p)
        elif line.startswith("@ "):
            org, _, dates = line[2:].partition("::")
            p = _tight(doc.add_paragraph(), before=6); right_tab(p)
            _run(p, org.strip(), bold=True)
            if dates.strip():
                _run(p, "\t" + dates.strip())
        elif line.startswith("> "):
            role, _, dates = line[2:].partition("::")
            p = _tight(doc.add_paragraph(), after=1); right_tab(p)
            _run(p, role.strip(), italic=True)
            if dates.strip():
                _run(p, "\t" + dates.strip(), italic=True)
        elif line.lstrip().startswith(("- ", "* ")):
            p = _tight(doc.add_paragraph(), after=2.5, line=1.05)
            p.paragraph_format.left_indent = Inches(0.23)
            p.paragraph_format.first_line_indent = Inches(-0.13)
            _run(p, "•  " + line.lstrip()[2:].strip())
        else:
            _run(_tight(doc.add_paragraph(), after=2, line=1.05), line.strip())
    return doc


def build_cover_docx(text: str):
    """Render the cover letter in Garamond, business-letter spacing, one page."""
    import docx
    from docx.shared import Inches

    doc = docx.Document()
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = GARAMOND, _Pt(11)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(1.0)
    for block in re.split(r"\n\s*\n", text.strip()):
        if not block.strip():
            continue
        p = _tight(doc.add_paragraph(), after=10, line=1.15)
        for i, ln in enumerate(block.split("\n")):
            if i:
                p.add_run().add_break()
            _run(p, ln.strip())
    return doc


def _docx_bytes(doc) -> bytes:
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_SLUG = re.compile(r"[^a-z0-9]+")


def _slug(s: str, n: int = 40) -> str:
    return _SLUG.sub("-", (s or "").lower()).strip("-")[:n] or "untitled"


def role_dir(company: str, title: str) -> Path:
    return config.APPLICATIONS_DIR / f"{_slug(company)}-{_slug(title)}"


def _save_local(conn, job, company, title, resume_md, cover_md, omitted, model) -> Dict[str, str]:
    d = role_dir(company, title)
    d.mkdir(parents=True, exist_ok=True)
    paths = {"resume_md": d / "resume.md", "resume_docx": d / "resume.docx",
             "cover_md": d / "cover_letter.md", "cover_docx": d / "cover_letter.docx"}
    resume_full = resume_md
    if omitted:  # honesty note as a markdown comment (kept out of the .docx body)
        resume_full += "\n\n<!-- JD requirements NOT claimed (no fabrication): " + "; ".join(map(str, omitted)) + " -->\n"
    paths["resume_md"].write_text(resume_full, encoding="utf-8")
    paths["cover_md"].write_text(cover_md, encoding="utf-8")
    build_resume_docx(resume_md).save(str(paths["resume_docx"]))
    build_cover_docx(cover_md).save(str(paths["cover_docx"]))
    store.record_draft(conn, job["id"], company=company, title=title, dir=str(d),
                       resume_md=paths["resume_md"], resume_docx=paths["resume_docx"],
                       cover_md=paths["cover_md"], cover_docx=paths["cover_docx"], model=model)
    return {"where": "local", "folder": str(d), "resume_url": str(paths["resume_docx"]),
            "cover_url": str(paths["cover_docx"])}


def generate_for_role(
    conn: sqlite3.Connection, job: sqlite3.Row, master: Dict[str, Any], voice: Dict[str, Any],
    *, model: str = config.DEEP_MODEL, regenerate: bool = False, to_drive: bool = True,
) -> Optional[Dict[str, str]]:
    """Generate + persist a draft set for one role. Writes editable Google Docs to the
    user's Drive (resume + cover letter), falling back to local files if Drive isn't
    available. Returns a dict of links, or None if already drafted (and not regenerate)."""
    if not regenerate and (store.get_draft(conn, job["id"])
                           or store.get_draft_for_role(conn, job["company"] or "", job["title"] or "")):
        return None  # never draft the same role twice (even under a re-fetched job id)

    full = store.get_job(conn, job["id"]) or job  # select_master rows omit `description`
    prompt = _draft_prompt(master, voice, full)
    text = llm.complete_text(prompt, model=model, system=DRAFT_SYSTEM, max_tokens=12000)
    resume_md, cover_md, omitted = _split_sections(text)
    if not resume_md or not cover_md:
        raise llm.LLMError(f"Draft generation returned incomplete content for job {job['id']}.")

    issues = _style_violations(cover_md)
    if issues:  # one corrective pass; the mechanical scrub below is the final guarantee
        retry = llm.complete_text(
            prompt + "\n\nYour previous attempt failed style review: " + "; ".join(issues)
            + ". Regenerate ALL sections with these problems fixed (same grounding rules).",
            model=model, system=DRAFT_SYSTEM, max_tokens=12000)
        r2, c2, o2 = _split_sections(retry)
        if r2 and c2:
            resume_md, cover_md, omitted = r2, c2, o2
    resume_md = _strip_md_bold(resume_md)
    cover_md = _scrub_dashes(_strip_md_bold(cover_md))

    company = job["company"] or "Company"
    title = job["title"] or "Role"

    if to_drive and oauth.is_authorized():
        # Signed in => drafts ALWAYS go to Drive. Retry a transient blip; on a real
        # failure surface a clear error (never silently save locally).
        err = None
        for _attempt in range(3):
            try:
                links = oauth.upload_drafts(conn, company, title,
                                            _docx_bytes(build_resume_docx(resume_md)),
                                            _docx_bytes(build_cover_docx(cover_md)))
                store.record_draft(conn, job["id"], company=company, title=title, dir=links["folder"],
                                   drive_url=links["folder"], resume_url=links["resume_url"],
                                   cover_url=links["cover_url"], model=model)
                return {"where": "drive", **links}
            except oauth.OAuthError as e:
                raise llm.LLMError(f"Google Drive sign-in expired — run `job-agent auth`, then retry. ({e})")
            except Exception as e:  # noqa: BLE001
                if "quota" in str(e).lower():  # account storage full — retrying won't help
                    raise llm.LLMError("Your Google account storage is full, so the draft can't be saved to "
                                       "Drive. Free up space (it's Gmail/Photos using it, not Drive) or "
                                       "upgrade your plan, then click Draft again.")
                err = e  # transient — retry
        raise llm.LLMError(f"Drive upload failed after 3 tries ({err}). Drafts are Drive-only — please retry.")
    if to_drive:  # not signed in at all -> local fallback so a draft is never lost
        print("   (Not signed in to Google Drive — run `job-agent auth`. Saving locally for now.)")

    return _save_local(conn, job, company, title, resume_md, cover_md, omitted, model)


def migrate_local_draft(conn: sqlite3.Connection, job: sqlite3.Row, draft: sqlite3.Row) -> Optional[Dict[str, str]]:
    """Upload an EXISTING local draft pair to Drive unchanged — no LLM call, so the
    user never gets a second, different pair. Returns links, or None if the local
    files are gone (caller may then regenerate)."""
    if not oauth.is_authorized():
        return None
    r, c = draft["resume_docx"], draft["cover_docx"]
    if not (r and c and Path(r).exists() and Path(c).exists()):
        return None
    company = draft["company"] or job["company"] or "Company"
    title = draft["title"] or job["title"] or "Role"
    links = oauth.upload_drafts(conn, company, title, Path(r).read_bytes(), Path(c).read_bytes())
    store.record_draft(conn, job["id"], company=company, title=title, dir=links["folder"],
                       drive_url=links["folder"], resume_url=links["resume_url"],
                       cover_url=links["cover_url"], model=draft["model"])
    return {"where": "drive", **links}


def run_drafts(
    conn: sqlite3.Connection, rows: List[sqlite3.Row], master: Dict[str, Any], voice: Dict[str, Any],
    *, model: str = config.DEEP_MODEL, regenerate: bool = False, to_drive: bool = True,
) -> Tuple[int, int]:
    """Generate drafts for the given rows. Returns (generated, skipped)."""
    generated = skipped = 0
    for job in rows:
        try:
            res = generate_for_role(conn, job, master, voice, model=model, regenerate=regenerate, to_drive=to_drive)
        except llm.LLMError as e:
            print(f"   ! draft failed for [{job['id']}] {job['company']} – {job['title']}: {e}")
            continue
        if res is None:
            skipped += 1
        else:
            generated += 1
            where = res.get("folder") if res.get("where") == "drive" else role_dir(job["company"] or "", job["title"] or "").name + "/"
            print(f"   + drafted: {job['company']} – {job['title']} -> {where}")
    return generated, skipped
