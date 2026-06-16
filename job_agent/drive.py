"""Google Drive (read-only) access — source material for the master profile.

Auth uses a service-account JSON (GOOGLE_SERVICE_ACCOUNT_JSON). A service account
only sees files explicitly shared with it, so if a search returns nothing the
caller STOPS and asks the user to share the folder — we never silently proceed or
invent content. We pull full text: Google Docs are exported as text/plain, PDFs go
through pypdf, .docx through python-docx. Returned text is used verbatim as grounding
for profile synthesis.

Network/Google libraries are imported lazily so the rest of the app (and its tests)
stay dependency-light.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import List, Optional, Tuple

from . import config

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

# Career documents we want (resumes / CVs / cover letters) vs. things that merely
# contain the candidate's name (tax forms, offers, essays, NDAs) which must NOT feed
# the profile. Filtering is by filename — controllable, and the final set is shown
# to the user to confirm coverage.
INCLUDE_RE = re.compile(r"resume|résumé|\bcv\b|cover|\bcl\b|curriculum", re.I)
EXCLUDE_RE = re.compile(r"\bw-?9\b|\bw-?2\b|\boffer\b|essay|1099|\btax\b|\bi-?9\b|\bnda\b", re.I)
COVER_RE = re.compile(r"cover|\bcl\b", re.I)

GOOGLE_DOC = "application/vnd.google-apps.document"
PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_LEGACY = "application/msword"


class DriveError(RuntimeError):
    """Raised for missing creds / disabled API / no shared files — caller must stop."""


def service_account_email() -> Optional[str]:
    """The SA email (for the 'please share your folder with…' message). No network."""
    path = config.GOOGLE_SERVICE_ACCOUNT_JSON
    if not path or not Path(path).exists():
        return None
    import json
    try:
        return json.loads(Path(path).read_text()).get("client_email")
    except Exception:
        return None


def build_service():
    """Authenticated Drive v3 client. Raises DriveError on missing/invalid creds."""
    path = config.GOOGLE_SERVICE_ACCOUNT_JSON
    if not path:
        raise DriveError("GOOGLE_SERVICE_ACCOUNT_JSON not set in .env.")
    if not Path(path).exists():
        raise DriveError(f"Service-account JSON not found at {path}.")
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
    except ImportError as e:  # pragma: no cover - dependency missing
        raise DriveError(f"Google libraries not installed: {e}") from e
    creds = service_account.Credentials.from_service_account_file(path, scopes=SCOPES)
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _list(svc, q: str) -> List[dict]:
    files, page = [], None
    while True:
        res = svc.files().list(
            q=q,
            fields="nextPageToken, files(id,name,mimeType,modifiedTime,size,owners(emailAddress))",
            pageSize=100, supportsAllDrives=True, includeItemsFromAllDrives=True,
            pageToken=page,
        ).execute()
        files.extend(res.get("files", []))
        page = res.get("nextPageToken")
        if not page:
            break
    return files


def find_documents(svc) -> List[dict]:
    """Resumes / CVs / cover letters shared with the SA, de-duped by id, newest first.

    Filters text-bearing docs (Doc/PDF/docx) by filename: INCLUDE résumé/CV/cover
    patterns, EXCLUDE tax forms / offers / essays / NDAs (which merely contain the
    candidate's name). Folders, sheets, slides, and images are ignored.
    """
    by_id = {}
    for f in _list(svc, "trashed = false"):
        if f.get("mimeType") not in (GOOGLE_DOC, PDF, DOCX, DOC_LEGACY):
            continue
        name = f.get("name", "")
        if EXCLUDE_RE.search(name):
            continue
        if INCLUDE_RE.search(name):
            by_id[f["id"]] = f
    return sorted(by_id.values(), key=lambda f: f.get("modifiedTime", ""), reverse=True)


def list_all_visible(svc) -> List[dict]:
    """Everything the SA can see — used to diagnose an empty result (sharing issue)."""
    return _list(svc, "trashed = false")


# --- text extraction --------------------------------------------------------

def _pdf_text(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join((p.extract_text() or "") for p in reader.pages).strip()


def _docx_text(raw: bytes) -> str:
    import docx  # python-docx
    d = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:  # cover letters rarely have tables, resumes sometimes do
        for row in table.rows:
            parts.append(" \t ".join(c.text for c in row.cells))
    return "\n".join(t for t in parts if t and t.strip()).strip()


def fetch_text(svc, f: dict) -> str:
    """Full text of one Drive file (export Doc as text; extract PDF/docx bytes)."""
    mt = f.get("mimeType")
    if mt == GOOGLE_DOC:
        data = svc.files().export(fileId=f["id"], mimeType="text/plain").execute()
        return (data.decode("utf-8", "replace") if isinstance(data, (bytes, bytearray)) else str(data)).strip()

    from googleapiclient.http import MediaIoBaseDownload
    buf = io.BytesIO()
    dl = MediaIoBaseDownload(buf, svc.files().get_media(fileId=f["id"], supportsAllDrives=True))
    done = False
    while not done:
        _, done = dl.next_chunk()
    raw = buf.getvalue()
    if mt == PDF:
        return _pdf_text(raw)
    if mt in (DOCX, DOC_LEGACY) or str(f.get("name", "")).lower().endswith(".docx"):
        return _docx_text(raw)
    try:
        return raw.decode("utf-8", "replace").strip()
    except Exception:
        return ""


def is_cover_letter(f: dict) -> bool:
    return bool(COVER_RE.search(str(f.get("name", ""))))


def collect() -> Tuple[List[dict], List[Tuple[dict, str]]]:
    """Authenticate, find docs, pull text. Returns (all_files_found, [(file, text)]).

    Raises DriveError if creds are missing/invalid. Returns empty lists (not an
    error) if nothing is shared — the caller decides to STOP and ask for sharing.
    """
    svc = build_service()
    files = find_documents(svc)
    docs: List[Tuple[dict, str]] = []
    for f in files:
        try:
            text = fetch_text(svc, f)
        except Exception:
            text = ""
        if text and len(text.strip()) >= 30:
            docs.append((f, text.strip()))
    return files, docs
